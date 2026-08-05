from __future__ import annotations

from datetime import datetime
from io import BytesIO
from pathlib import Path
import re
from zipfile import BadZipFile, ZipFile

from docx import Document as open_docx
from docx.document import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.oxml.ns import qn

from ..models.canonical_document import (
    CanonicalDocument,
    DocumentMetadata,
    ImageReference,
    Paragraph,
    Reference,
    Section,
    SourceProvenance,
    Table,
)
from .base import AbstractParser, InvalidParserInputError, normalize_mime_type


DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)


def _iso_datetime(value: datetime | None) -> str | None:
    return value.isoformat() if value is not None else None


class DocxParser(AbstractParser):
    """Deterministic Office Open XML text and structure extractor."""

    extensions = frozenset({".docx"})
    mime_types = frozenset({DOCX_MIME_TYPE})

    def parser_name(self) -> str:
        return "docx"

    def parser_version(self) -> str:
        return "2.0.0"

    def validate(self, content: bytes) -> bool:
        if not content:
            raise InvalidParserInputError("DOCX document cannot be empty")
        if not content.startswith(b"PK"):
            raise InvalidParserInputError(
                "DOCX document does not have a valid ZIP signature"
            )
        try:
            with ZipFile(BytesIO(content)) as archive:
                names = frozenset(archive.namelist())
                if "[Content_Types].xml" not in names:
                    raise InvalidParserInputError(
                        "DOCX package has no content type manifest"
                    )
                if "word/document.xml" not in names:
                    raise InvalidParserInputError(
                        "DOCX package has no main document part"
                    )
            self._document(content)
        except InvalidParserInputError:
            raise
        except (BadZipFile, PackageNotFoundError, KeyError, ValueError) as exc:
            raise InvalidParserInputError(
                "DOCX document is malformed or unreadable"
            ) from exc
        return True

    def extract_metadata(
        self,
        content: bytes,
        *,
        filename: str | None = None,
        mime_type: str | None = None,
    ) -> DocumentMetadata:
        self.validate(content)
        document = self._document(content)
        properties = document.core_properties
        title = self._clean(properties.title)
        if title is None and filename:
            title = self._filename_title(filename)
        author = self._clean(properties.author)
        keywords = [
            value.strip()
            for value in re.split(r"[,;]", properties.keywords or "")
            if value.strip()
        ]
        page_break_count = len(
            document.element.body.xpath(
                './/w:br[@w:type="page"] | .//w:lastRenderedPageBreak'
            )
        )
        return DocumentMetadata(
            title=title,
            author=author,
            authors=[author] if author else [],
            subject=self._clean(properties.subject),
            keywords=keywords,
            creator=self._clean(properties.last_modified_by),
            created_at=_iso_datetime(properties.created),
            modified_at=_iso_datetime(properties.modified),
            language=self._language(document),
            source_filename=Path(filename).name if filename else None,
            extension=".docx",
            mime_type=normalize_mime_type(mime_type) or DOCX_MIME_TYPE,
            page_count=page_break_count + 1,
            properties={
                "category": self._clean(properties.category),
                "comments": self._clean(properties.comments),
                "content_status": self._clean(properties.content_status),
                "identifier": self._clean(properties.identifier),
                "last_modified_by": self._clean(
                    properties.last_modified_by
                ),
                "revision": properties.revision,
                "page_break_count": page_break_count,
            },
        )

    def parse(
        self,
        content: bytes,
        *,
        filename: str | None = None,
        mime_type: str | None = None,
    ) -> CanonicalDocument:
        self.validate(content)
        document = self._document(content)
        metadata = self.extract_metadata(
            content,
            filename=filename,
            mime_type=mime_type,
        )
        numbering = self._numbering_formats(document)
        sections: list[Section] = []
        section_stack: list[Section] = []
        tables: list[Table] = []
        references: list[Reference] = []
        images: list[ImageReference] = []
        content_blocks: list[str] = []
        source_index = 0
        paragraph_index = 0
        table_index = 0
        page_number = 1

        def section_path() -> list[str]:
            return [
                section.title
                for section in section_stack
                if section.title is not None
            ]

        def add_section(title: str, level: int, index: int) -> None:
            nonlocal section_stack
            while section_stack and section_stack[-1].level >= level:
                section_stack.pop()
            section = Section(
                title=title,
                level=level,
                order=index,
                page_start=page_number,
                page_end=page_number,
                provenance=SourceProvenance(
                    source_index=index,
                    page_number=page_number,
                    section_path=section_path() + [title],
                    source_locator=f"word/body[{index}]",
                ),
            )
            if section_stack:
                section_stack[-1].subsections.append(section)
            else:
                sections.append(section)
            section_stack.append(section)

        def ensure_section(index: int) -> Section:
            if not section_stack:
                title = metadata.title or "Document"
                add_section(title, 1, index)
            return section_stack[-1]

        for child in document.element.body.iterchildren():
            if child.tag == qn("w:p"):
                docx_paragraph = DocxParagraph(child, document)
                text, page_breaks = self._paragraph_text(docx_paragraph)
                heading_level = self._heading_level(docx_paragraph)
                if heading_level is not None and text:
                    clean_heading = self._normalize_block_text(text)
                    add_section(clean_heading, heading_level, source_index)
                    content_blocks.append(clean_heading)
                    references.extend(
                        self._hyperlinks(
                            docx_paragraph,
                            source_index=source_index,
                            page_number=page_number,
                            section_path=section_path(),
                        )
                    )
                    images.extend(
                        self._images(
                            child,
                            document,
                            source_index=source_index,
                            page_number=page_number,
                            section_path=section_path(),
                        )
                    )
                elif text.strip():
                    current = ensure_section(source_index)
                    path = section_path()
                    list_metadata = self._list_metadata(
                        docx_paragraph,
                        numbering,
                    )
                    clean_text = self._normalize_block_text(text)
                    start = self._character_start(content_blocks)
                    paragraph = Paragraph(
                        text=clean_text,
                        order=paragraph_index,
                        page_number=page_number,
                        provenance=SourceProvenance(
                            source_index=source_index,
                            page_number=page_number,
                            section_path=path,
                            paragraph_index=paragraph_index,
                            character_start=start,
                            character_end=start + len(clean_text),
                            source_locator=f"word/body[{source_index}]/p",
                        ),
                        metadata={
                            "style": (
                                docx_paragraph.style.name
                                if docx_paragraph.style is not None
                                else None
                            ),
                            "page_break_count": page_breaks,
                            **list_metadata,
                        },
                    )
                    current.paragraphs.append(paragraph)
                    current.page_end = page_number + page_breaks
                    references.extend(
                        self._hyperlinks(
                            docx_paragraph,
                            source_index=source_index,
                            page_number=page_number,
                            section_path=path,
                        )
                    )
                    images.extend(
                        self._images(
                            child,
                            document,
                            source_index=source_index,
                            page_number=page_number,
                            section_path=path,
                        )
                    )
                    content_blocks.append(clean_text)
                    paragraph_index += 1
                else:
                    images.extend(
                        self._images(
                            child,
                            document,
                            source_index=source_index,
                            page_number=page_number,
                            section_path=section_path(),
                        )
                    )
                page_number += page_breaks
            elif child.tag == qn("w:tbl"):
                docx_table = DocxTable(child, document)
                rows = [
                    [
                        self._normalize_cell_text(cell.text)
                        for cell in row.cells
                    ]
                    for row in docx_table.rows
                ]
                if rows:
                    path = section_path()
                    headers = rows[0]
                    body_rows = rows[1:]
                    table = Table(
                        headers=headers,
                        rows=body_rows,
                        page_number=page_number,
                        provenance=SourceProvenance(
                            source_index=source_index,
                            page_number=page_number,
                            section_path=path,
                            table_index=table_index,
                            source_locator=f"word/body[{source_index}]/tbl",
                        ),
                        metadata={
                            "row_count": len(rows),
                            "column_count": max(
                                (len(row) for row in rows),
                                default=0,
                            ),
                        },
                    )
                    tables.append(table)
                    flattened = "\n".join(
                        "\t".join(cell for cell in row)
                        for row in rows
                        if any(cell for cell in row)
                    ).strip()
                    if flattened:
                        content_blocks.append(flattened)
                    table_index += 1
                for row in docx_table.rows:
                    for cell in row.cells:
                        for cell_paragraph in cell.paragraphs:
                            references.extend(
                                self._hyperlinks(
                                    cell_paragraph,
                                    source_index=source_index,
                                    page_number=page_number,
                                    section_path=section_path(),
                                )
                            )
                images.extend(
                    self._images(
                        child,
                        document,
                        source_index=source_index,
                        page_number=page_number,
                        section_path=section_path(),
                    )
                )
            source_index += 1

        if not any(True for section in sections for _ in section.iter_paragraphs()):
            flattened = "\n\n".join(content_blocks).strip()
            if not flattened:
                raise InvalidParserInputError(
                    "DOCX document contains no extractable text"
                )
            root = ensure_section(0)
            root.paragraphs.append(
                Paragraph(
                    text=flattened,
                    order=0,
                    page_number=1,
                    provenance=SourceProvenance(
                        source_index=0,
                        page_number=1,
                        section_path=[root.title] if root.title else [],
                        paragraph_index=0,
                        character_start=0,
                        character_end=len(flattened),
                        source_locator="word/document.xml",
                    ),
                    metadata={"block_type": "structural_fallback"},
                )
            )
        canonical = CanonicalDocument.build(
            parser_name=self.parser_name(),
            parser_version=self.parser_version(),
            metadata=metadata,
            content="\n\n".join(content_blocks),
            sections=sections,
            tables=tables,
            images=self._deduplicate_images(images),
            references=references,
        )
        return canonical

    @staticmethod
    def _document(content: bytes) -> DocxDocument:
        try:
            return open_docx(BytesIO(content))
        except (BadZipFile, PackageNotFoundError, KeyError, ValueError) as exc:
            raise InvalidParserInputError(
                "DOCX document is malformed or unreadable"
            ) from exc

    @staticmethod
    def _clean(value: object | None) -> str | None:
        if value is None:
            return None
        normalized = str(value).strip()
        return normalized or None

    @staticmethod
    def _filename_title(filename: str) -> str | None:
        value = Path(filename).stem.replace("_", " ").replace("-", " ")
        value = re.sub(r"\s+", " ", value).strip()
        return value or None

    @staticmethod
    def _language(document: DocxDocument) -> str | None:
        for language in document.element.body.xpath(".//w:lang"):
            value = language.get(qn("w:val"))
            if value:
                return value.strip() or None
        return None

    @staticmethod
    def _heading_level(paragraph: DocxParagraph) -> int | None:
        style_name = (
            paragraph.style.name.strip()
            if paragraph.style is not None and paragraph.style.name
            else ""
        )
        match = re.match(r"^Heading\s+([1-9]\d*)$", style_name, re.I)
        if match:
            return min(int(match.group(1)), 9)
        outline = paragraph._p.pPr
        if outline is not None and outline.outlineLvl is not None:
            value = outline.outlineLvl.val
            try:
                return min(int(value) + 1, 9)
            except (TypeError, ValueError):
                return None
        return None

    @staticmethod
    def _paragraph_text(paragraph: DocxParagraph) -> tuple[str, int]:
        parts: list[str] = []
        page_breaks = 0
        for node in paragraph._p.iter():
            if node.tag == qn("w:t") and node.text:
                parts.append(node.text)
            elif node.tag == qn("w:tab"):
                parts.append("\t")
            elif node.tag == qn("w:br"):
                if node.get(qn("w:type")) == "page":
                    parts.append("\n[PAGE_BREAK]\n")
                    page_breaks += 1
                else:
                    parts.append("\n")
            elif node.tag == qn("w:lastRenderedPageBreak"):
                parts.append("\n[PAGE_BREAK]\n")
                page_breaks += 1
        return "".join(parts), page_breaks

    @staticmethod
    def _normalize_block_text(value: str) -> str:
        lines = [
            re.sub(r"[ \t]+", " ", line).strip()
            for line in value.replace("\r", "\n").splitlines()
        ]
        return "\n".join(line for line in lines if line).strip()

    @staticmethod
    def _normalize_cell_text(value: str) -> str:
        return re.sub(r"\s+", " ", value).strip()

    @staticmethod
    def _character_start(content_blocks: list[str]) -> int:
        return sum(len(block) for block in content_blocks) + 2 * len(
            content_blocks
        )

    @staticmethod
    def _numbering_formats(
        document: DocxDocument,
    ) -> dict[tuple[str, int], str]:
        try:
            numbering = document.part.numbering_part.element
        except (AttributeError, KeyError):
            return {}
        abstract_formats: dict[tuple[str, int], str] = {}
        for abstract in numbering.findall(qn("w:abstractNum")):
            abstract_id = abstract.get(qn("w:abstractNumId"))
            if abstract_id is None:
                continue
            for level in abstract.findall(qn("w:lvl")):
                raw_level = level.get(qn("w:ilvl"))
                num_format = level.find(qn("w:numFmt"))
                if raw_level is None or num_format is None:
                    continue
                value = num_format.get(qn("w:val"))
                if value:
                    abstract_formats[(abstract_id, int(raw_level))] = value
        result: dict[tuple[str, int], str] = {}
        for number in numbering.findall(qn("w:num")):
            number_id = number.get(qn("w:numId"))
            abstract_ref = number.find(qn("w:abstractNumId"))
            if number_id is None or abstract_ref is None:
                continue
            abstract_id = abstract_ref.get(qn("w:val"))
            if abstract_id is None:
                continue
            for (candidate_id, level), value in abstract_formats.items():
                if candidate_id == abstract_id:
                    result[(number_id, level)] = value
        return result

    @staticmethod
    def _list_metadata(
        paragraph: DocxParagraph,
        numbering: dict[tuple[str, int], str],
    ) -> dict[str, object]:
        properties = paragraph._p.pPr
        style_name = (
            paragraph.style.name.casefold()
            if paragraph.style is not None and paragraph.style.name
            else ""
        )
        if properties is None or properties.numPr is None:
            if style_name.startswith("list bullet"):
                level_match = re.search(r"(\d+)$", style_name)
                return {
                    "block_type": "list_item",
                    "list_type": "unordered",
                    "list_level": (
                        max(0, int(level_match.group(1)) - 1)
                        if level_match
                        else 0
                    ),
                    "list_id": None,
                    "number_format": "bullet",
                }
            if style_name.startswith("list number"):
                level_match = re.search(r"(\d+)$", style_name)
                return {
                    "block_type": "list_item",
                    "list_type": "ordered",
                    "list_level": (
                        max(0, int(level_match.group(1)) - 1)
                        if level_match
                        else 0
                    ),
                    "list_id": None,
                    "number_format": "decimal",
                }
            return {"block_type": "paragraph"}
        num_id_node = properties.numPr.numId
        level_node = properties.numPr.ilvl
        if num_id_node is None:
            return {"block_type": "paragraph"}
        num_id = str(num_id_node.val)
        level = int(level_node.val) if level_node is not None else 0
        number_format = numbering.get((num_id, level), "decimal")
        unordered = number_format in {
            "bullet",
            "none",
        }
        return {
            "block_type": "list_item",
            "list_type": "unordered" if unordered else "ordered",
            "list_level": level,
            "list_id": num_id,
            "number_format": number_format,
        }

    @staticmethod
    def _hyperlinks(
        paragraph: DocxParagraph,
        *,
        source_index: int,
        page_number: int,
        section_path: list[str],
    ) -> list[Reference]:
        references: list[Reference] = []
        for link_index, hyperlink in enumerate(
            paragraph._p.findall(".//" + qn("w:hyperlink"))
        ):
            text = "".join(
                node.text or ""
                for node in hyperlink.iter(qn("w:t"))
            ).strip()
            relation_id = hyperlink.get(qn("r:id"))
            anchor = hyperlink.get(qn("w:anchor"))
            target = None
            if relation_id and relation_id in paragraph.part.rels:
                target = paragraph.part.rels[relation_id].target_ref
            elif anchor:
                target = f"#{anchor}"
            if text or target:
                references.append(
                    Reference(
                        text=text or target or "hyperlink",
                        target=target,
                        reference_type="hyperlink",
                        page_number=page_number,
                        provenance=SourceProvenance(
                            source_index=source_index,
                            page_number=page_number,
                            section_path=section_path,
                            source_locator=(
                                f"word/body[{source_index}]/"
                                f"hyperlink[{link_index}]"
                            ),
                        ),
                        metadata={"relationship_id": relation_id},
                    )
                )
        return references

    @staticmethod
    def _images(
        element,
        document: DocxDocument,
        *,
        source_index: int,
        page_number: int,
        section_path: list[str],
    ) -> list[ImageReference]:
        images: list[ImageReference] = []
        descriptions = [
            node.get("descr") or node.get("title") or node.get("name")
            for node in element.findall(".//" + qn("wp:docPr"))
        ]
        for image_index, blip in enumerate(
            element.findall(".//" + qn("a:blip"))
        ):
            relation_id = blip.get(qn("r:embed"))
            if not relation_id:
                continue
            relationship = document.part.rels.get(relation_id)
            target = relationship.target_ref if relationship else None
            target_part = relationship.target_part if relationship else None
            alt_text = (
                descriptions[image_index]
                if image_index < len(descriptions)
                else None
            )
            images.append(
                ImageReference(
                    identifier=relation_id,
                    source=target,
                    alt_text=alt_text,
                    page_number=page_number,
                    provenance=SourceProvenance(
                        source_index=source_index,
                        page_number=page_number,
                        section_path=section_path,
                        source_locator=(
                            f"word/body[{source_index}]/"
                            f"drawing[{image_index}]"
                        ),
                    ),
                    metadata={
                        "content_type": getattr(
                            target_part,
                            "content_type",
                            None,
                        ),
                        "embedded": True,
                        "ocr_performed": False,
                    },
                )
            )
        return images

    @staticmethod
    def _deduplicate_images(
        images: list[ImageReference],
    ) -> list[ImageReference]:
        seen: set[tuple[str, int | None, int]] = set()
        result: list[ImageReference] = []
        for image in images:
            source_index = (
                image.provenance.source_index
                if image.provenance is not None
                else -1
            )
            key = (image.identifier, image.page_number, source_index)
            if key not in seen:
                seen.add(key)
                result.append(image)
        return result


DOCXParser = DocxParser


__all__ = ["DOCXParser", "DocxParser"]
