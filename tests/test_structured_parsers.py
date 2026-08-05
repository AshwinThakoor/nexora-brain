from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
import base64

import pytest
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

from nexora_knowledge.parsers import DocxParser, HtmlParser, MarkdownParser


def _add_hyperlink(paragraph, text: str, target: str) -> None:
    relation_id = paragraph.part.relate_to(
        target,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _docx_bytes() -> bytes:
    document = Document()
    document.core_properties.title = "Deterministic DOCX"
    document.core_properties.author = "NEXORA"
    document.core_properties.created = datetime(
        2026,
        7,
        28,
        tzinfo=timezone.utc,
    )
    document.add_heading("Introduction", level=1)
    paragraph = document.add_paragraph("Visit ")
    _add_hyperlink(paragraph, "Example", "https://example.invalid/docs")
    document.add_paragraph("First bullet", style="List Bullet")
    document.add_paragraph("First numbered", style="List Number")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "alpha"
    table.cell(1, 1).text = "1"
    page_paragraph = document.add_paragraph("Before")
    page_paragraph.add_run().add_break(WD_BREAK.PAGE)
    page_paragraph.add_run("After")
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0"
        "lEQVR42mP8/x8AAusB9Y9Z9ZsAAAAASUVORK5CYII="
    )
    document.add_picture(BytesIO(png))
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_docx_parser_extracts_properties_structure_and_references() -> None:
    parsed = DocxParser().parse(
        _docx_bytes(),
        filename="fixture.docx",
        mime_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    assert parsed.metadata.title == "Deterministic DOCX"
    assert parsed.metadata.author == "NEXORA"
    assert parsed.metadata.created_at.startswith("2026-07-28")
    assert parsed.metadata.page_count == 2
    assert parsed.sections[0].title == "Introduction"
    paragraphs = list(parsed.iter_paragraphs())
    assert any(item.text == "Visit Example" for item in paragraphs)
    assert any(
        item.metadata.get("list_type") == "unordered"
        for item in paragraphs
    )
    assert any(
        item.metadata.get("list_type") == "ordered"
        for item in paragraphs
    )
    assert parsed.tables[0].headers == ["Name", "Value"]
    assert parsed.tables[0].rows == [["alpha", "1"]]
    assert parsed.references[0].target == "https://example.invalid/docs"
    assert parsed.images
    assert parsed.images[0].metadata["ocr_performed"] is False
    assert "[PAGE_BREAK]" in parsed.content
    parsed.assert_valid()


def test_docx_parser_rejects_malformed_package() -> None:
    with pytest.raises(ValueError, match="DOCX"):
        DocxParser().parse(b"PK-not-a-real-package", filename="bad.docx")


def test_markdown_parser_extracts_all_deterministic_blocks(tmp_path) -> None:
    marker = tmp_path / "must-not-exist"
    source = f"""---
title: Markdown Fixture
author: NEXORA
keywords: [parser, provenance]
---
# Overview

Paragraph with `inline_code()` and [link](https://example.invalid).

- bullet
1. numbered

> quoted content

```python
open({str(marker)!r}, "w").write("unsafe")
```

| Name | Value |
| --- | ---: |
| alpha | 1 |

![diagram](https://example.invalid/image.png)

---

[manual]: https://example.invalid/manual "Manual"
"""
    parsed = MarkdownParser().parse(
        source.encode(),
        filename="fixture.md",
        mime_type="text/markdown",
    )

    assert parsed.metadata.title == "Markdown Fixture"
    assert parsed.metadata.author == "NEXORA"
    assert parsed.metadata.properties["front_matter"]["keywords"] == [
        "parser",
        "provenance",
    ]
    assert parsed.sections[0].title == "Overview"
    paragraphs = list(parsed.iter_paragraphs())
    assert any(item.metadata["inline_code"] for item in paragraphs)
    assert any(
        item.metadata.get("list_type") == "unordered"
        for item in paragraphs
    )
    assert any(
        item.metadata.get("list_type") == "ordered"
        for item in paragraphs
    )
    assert any(
        item.metadata.get("block_type") == "blockquote"
        for item in paragraphs
    )
    assert any(
        item.metadata.get("block_type") == "fenced_code"
        and item.metadata["executed"] is False
        for item in paragraphs
    )
    assert parsed.tables[0].headers == ["Name", "Value"]
    assert any(item.reference_type == "link" for item in parsed.references)
    assert any(
        item.reference_type == "reference_definition"
        for item in parsed.references
    )
    assert parsed.images[0].metadata["loaded"] is False
    assert not marker.exists()
    parsed.assert_valid()


def test_html_parser_extracts_content_and_ignores_unsafe_nodes() -> None:
    source = b"""<!doctype html>
<html lang="en"><head>
  <title>HTML Fixture</title>
  <meta name="author" content="NEXORA">
  <meta name="description" content="Deterministic HTML">
  <style>.hidden { display:none }</style>
  <script>document.body.textContent = 'executed';</script>
</head><body>
  <h1>Overview</h1>
  <p>Paragraph with <code>inline()</code> and
     <a href="https://example.invalid/page">a link</a>.</p>
  <ul><li>bullet</li></ul>
  <ol><li>numbered</li></ol>
  <blockquote>quoted</blockquote>
  <pre><code>print("not executed")</code></pre>
  <table><tr><th>Name</th><th>Value</th></tr>
         <tr><td>alpha</td><td>1</td></tr></table>
  <img src="https://example.invalid/image.png" alt="diagram">
  <iframe src="https://example.invalid/frame"></iframe>
</body></html>"""
    parsed = HtmlParser().parse(
        source,
        filename="fixture.html",
        mime_type="text/html",
    )

    assert parsed.metadata.title == "HTML Fixture"
    assert parsed.metadata.author == "NEXORA"
    assert parsed.metadata.language == "en"
    assert parsed.metadata.properties["external_resources_loaded"] is False
    assert "document.body" not in parsed.content
    assert "display:none" not in parsed.content
    assert parsed.sections[0].title == "Overview"
    paragraphs = list(parsed.iter_paragraphs())
    assert any(item.metadata["inline_code"] for item in paragraphs)
    assert {item.metadata.get("list_type") for item in paragraphs} >= {
        "ordered",
        "unordered",
    }
    assert any(
        item.metadata.get("block_type") == "preformatted_code"
        for item in paragraphs
    )
    assert parsed.tables[0].rows == [["alpha", "1"]]
    assert parsed.references[0].metadata["loaded"] is False
    assert parsed.images[0].metadata["loaded"] is False
    parsed.assert_valid()


def test_html_parser_handles_malformed_markup_and_enforces_limits() -> None:
    malformed = HtmlParser().parse(
        b"<html><body><h1>Title<p>Still content",
        filename="malformed.html",
        mime_type="text/html",
    )
    assert "Still content" in malformed.content

    limited = HtmlParser()
    limited.max_document_bytes = 10
    with pytest.raises(ValueError, match="size limit"):
        limited.parse(b"<p>This is larger than ten bytes</p>")
